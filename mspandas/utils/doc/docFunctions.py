"""Convenience tools for accessing doc templates using python-docx.
"""

import pandas as pd

import docx

# Microsoft cell margins in inches
margins_master = {
    'normal': {'top': 0.05, 'bottom': 0.05, 'left': 0.1, 'right': 0.1},
	'none': {'top': 0, 'bottom': 0, 'left': 0, 'right': 0},
	'narrow': {'top': 0.05, 'bottom': 0.05, 'left': 0.05, 'right': 0.05},
	'wide': {'top': 0.15, 'bottom': 0.15, 'left': 0.15, 'right': 0.15},
	# custom style, does not exist in doc
	'tight': {'top': 0.025, 'bottom': 0.025, 'left': 0.025, 'right': 0.025},
}

def set_row_height(table, row_height=0.15):
    """Set row height for table rows.

    Parameters
    ----------
    row_height: float
        Row height in inches.

    Returns:
    --------
    table: docx.table.Table
        docx table graphic frame object, with new row height.
    """
    emu = row_height * docx.shared.Length._EMUS_PER_INCH
    for r in table.rows:
        r.height = docx.shared.Emu(round(emu))
    return table

def format_cell(cell,
                font_size=None,
                font_color=None,
                font_name=None,
                bold=False,
                fill=False,
                fill_color=None,
                cell_margins='tight'):
    """Format a given table cell.

    Parameters
    ----------
    font_size: int
        Cell text font size. For more, see docx.shared.Pt
    font_color: tuple
        Cell text font color. Must be RGB code as tuple of 3 integers, or HEX code as string. For more see docx.shared.RGBColor
    font_name:
        Cell text font name, for example 'Arial'.
    bold: bool
        Whether or not to bold the text.
    fill: bool
        Whether or not to fill the cell backgound color.
    fill_color: tuple or docx.enum.dml.MSO_THEME_COLOR
        Color to fill cell background. Must be RGB code as tuple of 3 integers, or instance of docx.enum.dml.MSO_THEME_COLOR.
    cell_margins: str
        Keyword for setting cell margin widths. Use one of 'normal', 'none', 'narrow', 'tight', or 'wide'. Keywords are adopted from doc with a custom tight setting.

    Returns:
    --------
    cell: docx.table._Cell
        Table cell with applied formatting.
    """
    cell.margin_top = docx.shared.Inches(margins_master[cell_margins]['top'])
    cell.margin_bottom = docx.shared.Inches(margins_master[cell_margins]['bottom'])
    cell.margin_left = docx.shared.Inches(margins_master[cell_margins]['left'])
    cell.margin_right = docx.shared.Inches(margins_master[cell_margins]['right'])

    if fill:
        if isinstance(fill_color,docx.enum.base.EnumValue):
            xml_shd = docx.oxml.parse_xml(r'<w:shd {} w:fill="{}"/>'.format(docx.oxml.ns.nsdecls('w'), fill_color))
            cell._tc.get_or_add_tcPr().append(xml_shd)
        elif isinstance(fill_color,tuple):
            rgb = docx.shared.RGBColor(*fill_color)
            xml_shd = docx.oxml.parse_xml(r'<w:shd {} w:fill="{}"/>'.format(docx.oxml.ns.nsdecls('w'), rgb))
            cell._tc.get_or_add_tcPr().append(xml_shd)
        elif isinstance(fill_color,str):
            rgb = docx.shared.RGBColor.from_string(fill_color) if not fill_color.startswith('#') else docx.shared.RGBColor.from_string(fill_color[1:])
            xml_shd = docx.oxml.parse_xml(r'<w:shd {} w:fill="{}"/>'.format(docx.oxml.ns.nsdecls('w'), rgb))
            cell._tc.get_or_add_tcPr().append(xml_shd)
        else:
            raise ValueError('Incorrect value for fill_color. \
            Please provide one of RGB code as `tuple` of 3 integers or HEX code as string')

    try:
        p = cell.paragraphs[0]
    except IndexError:
        p = cell.add_paragraph()
    try:
        r = p.runs[0]
    except IndexError:
        r = p.add_run()
    if not font_size == None:
        r.font.size = docx.shared.Pt(font_size)
    if not font_color == None:
        if isinstance(font_color,tuple):
            r.font.color.rgb = docx.shared.RGBColor(*font_color)
        elif isinstance(font_color,str):
            r.font.color.rgb = docx.shared.RGBColor.from_string(font_color[1:]) if font_color.startswith('#') else docx.shared.RGBColor.from_string(font_color)
        else:
            raise ValueError('Incorrect value for font_color. \
            Please provide one of RGB code as `tuple` of 3 integers, or a HEX code as string')
    if not font_name == None:
        r.font.name = font_name
    if bold:
        r.font.bold = True
    return cell
